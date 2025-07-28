from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import google.generativeai as genai
from dotenv import load_dotenv
import tempfile
import PyPDF2  # ✅ required to read PDFs
from docx import Document # ✅ required to read DOCX files
from PIL import Image
import pytesseract  # ✅ required for OCR on images

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

app = Flask(__name__)
CORS(app)

# 🧠 Gemini Query Function
def ask_gemini(prompt):
    try:
        model = genai.GenerativeModel('models/gemini-2.5-pro')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print("Gemini Error:", e)
        return "An error occurred while querying Gemini."

# ✨ Route: Ask Gemini via GET/POST prompt
@app.route("/ask", methods=["GET", "POST"])
def ask():
    if request.method == "GET":
        return '''
            <form method="post" enctype="multipart/form-data" action="/ask">
                <input name="prompt" placeholder="Ask something" style="width:300px;" />
                <input type="file" name="file" />
                <input type="submit" />
            </form>
        '''
    try:
        prompt = request.form.get("prompt", "") or request.json.get("prompt", "")
        file = request.files.get("file")

        content = ""

        if file:
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                file.save(tmp.name)
                file_path = tmp.name

            if file.filename.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        content += page.extract_text()
            elif file.filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file.filename.endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
            elif file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                import pytesseract
                from PIL import Image
                image = Image.open(file_path)
                content = pytesseract.image_to_string(image)
            else:
                return jsonify({"error": "Unsupported file type."}), 400

        # Combine file content and prompt
        full_prompt = f"{prompt.strip()}\n\n{content[:15000]}" if content else prompt.strip()
        response = ask_gemini(full_prompt)
        return jsonify({"response": response})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ✅ Route: List Available Gemini Models
@app.route("/models", methods=["GET"])
def list_models():
    try:
        models = genai.list_models()
        return jsonify({"models": [m.name for m in models]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# 📄 NEW ROUTE: Summarize an uploaded document (PDF/Text)
# Add this if Tesseract is not in PATH (customize for your system)

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

@app.route("/summarize", methods=["POST"])
def summarize_document():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        filename = file.filename.lower()

        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            file.save(tmp.name)
            file_path = tmp.name

        content = ""

        # PDF
        if filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    content += page.extract_text() or ""

        # TXT
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

        # DOCX
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"

        # Image
        elif filename.endswith(('.png', '.jpg', '.jpeg')):
            image = Image.open(file_path)
            content = pytesseract.image_to_string(image)

        else:
            return jsonify({"error": "Unsupported file format. Upload PDF, TXT, DOCX, or image."}), 400

        if not content.strip():
            return jsonify({"error": "File is empty or unreadable."}), 400

        # Gemini summary
        prompt = f"Summarize the following document:\n\n{content[:15000]}"
        summary = ask_gemini(prompt)

        return jsonify({"summary": summary})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)
